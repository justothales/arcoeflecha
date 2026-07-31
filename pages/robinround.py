import io
import math
from pathlib import Path

import pandas as pd
import streamlit as st
import os
import shutil
import tempfile
from docx import Document
from PyPDF2 import PdfMerger
try:
    from docx2pdf import convert as docx2pdf_convert
except Exception:
    docx2pdf_convert = None

def to_excel(dfs, multi_sheet=False):
    """
    Converts a pandas DataFrame (or a dictionary of DataFrames for multi-sheet)
    into an in-memory Excel file (bytes).
    """
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        if multi_sheet:
            for sheet_name, df in dfs.items():
                df.to_excel(writer, sheet_name=sheet_name, index=False)
        else:
            dfs.to_excel(writer, index=False, sheet_name='Sheet1')
    return output.getvalue()


def load_dist_grupos():
    path = Path(__file__).resolve().parents[1] / 'DistGrupos.xlsx'
    if not path.exists():
        raise FileNotFoundError(f'Arquivo de distribuiÃ§Ã£o de grupos nÃ£o encontrado: {path}')
    df = pd.read_excel(path, index_col=0)
    df.index = pd.to_numeric(df.index, errors='coerce').astype('Int64')
    return df


def load_categoria_map():
    path = Path(__file__).resolve().parents[1] / 'DePara Categorias.xlsx'
    if not path.exists():
        raise FileNotFoundError(f'Arquivo de mapeamento de categorias nÃ£o encontrado: {path}')
    return pd.read_excel(path)


def load_input_df(uploaded_file):
    filename = getattr(uploaded_file, 'name', '')
    if filename.lower().endswith(('.txt', '.csv')):
        return pd.read_csv(uploaded_file, sep=';', dtype=str, engine='python')
    return pd.read_excel(uploaded_file, dtype=str)


def normalize_input(df, categoria_map):
    df = df.copy()
    # Build Nome Completo from existing columns or from FamilyName + GivenName
    if 'Nome Completo' in df.columns:
        df['Nome Completo'] = df['Nome Completo'].fillna('').astype(str).str.strip()
    elif 'NomeCompleto' in df.columns:
        df['Nome Completo'] = df['NomeCompleto'].fillna('').astype(str).str.strip()
    else:
        df['FamilyName'] = df.get('FamilyName', '').fillna('')
        df['GivenName'] = df.get('GivenName', '').fillna('')
        df['Nome Completo'] = (df['FamilyName'].astype(str).str.strip() + ' ' + df['GivenName'].astype(str).str.strip()).str.strip()

    df['Categoria Quali'] = (
        df.get('Division', '').fillna('').astype(str).str.strip()
        + df.get('Class', '').fillna('').astype(str).str.strip()
    ).str.strip()

    df = df.rename(columns={'Noc': 'Sigla', 'Country': 'Clube'})

    for col in ['Score', '10+X', 'X']:
        if col in df.columns:
            df[col] = pd.to_numeric(
                df[col].astype(str).str.replace(',', '.', regex=False).str.strip(),
                errors='coerce'
            ).fillna(0).astype(int)
        else:
            df[col] = 0

    categoria_lookup = categoria_map.set_index('Categoria Quali')['Categoria Combates']
    df['Categoria Combates'] = df['Categoria Quali'].map(categoria_lookup).fillna(df['Categoria Quali'])

    if 'WaID' not in df.columns:
        df['WaID'] = df.get('WaID', pd.NA)

    df['Clube'] = df.get('Clube', '').fillna('')
    df['Sigla'] = df.get('Sigla', '').fillna('')

    return df


def calculate_rank(df):
    df = df.copy()
    df = df.sort_values(
        by=['Categoria Combates', 'Score', '10+X', 'X', 'Nome Completo'],
        ascending=[True, False, False, False, True],
        kind='mergesort'
    )
    df['Rank'] = df.groupby('Categoria Combates').cumcount() + 1
    return df


def render_selection_table(df):
    try:
        if hasattr(st, 'data_editor'):
            return st.data_editor(df, hide_index=True, use_container_width=True)
        return st.experimental_data_editor(df, hide_index=True, use_container_width=True)
    except Exception:
        st.warning('O editor de dados nÃ£o estÃ¡ disponÃ­vel. Use os checkboxes abaixo para selecionar os atletas.')
        selected_indexes = []
        st.write('### Selecione os atletas que seguem para a prÃ³xima fase')
        for idx, row in df.iterrows():
            cols = st.columns([0.6, 2.2, 3.0, 1.5])
            selected = cols[0].checkbox('Sim', value=True, key=f'sel_{idx}')
            cols[1].write(row['Alvo'])
            cols[2].write(row['Nome Completo'])
            cols[3].write(row['Categoria Quali'])
            if selected:
                selected_indexes.append(idx)
        return df.loc[selected_indexes].copy()


def processar_combates(df, df_dist_grupos, etapa, local_da_prova):
    """
    Core logic to process the combat data.
    Takes dataframes and user inputs, returns three dataframes for output.
    """
    df = df.copy()
    df = df[df['Nome Completo'].astype(str).str.strip() != '']
    df = df[df['Categoria Combates'].astype(str).str.strip() != '']

    tabelas_division = {}
    for categoria_combate, data in df.groupby('Categoria Combates'):
        tabelas_division[categoria_combate] = data.copy()

    tabela_eliminados = pd.DataFrame()
    combates_geral_df = pd.DataFrame()
    final_grupos_tabelas = {}

    for categoria_combate, tabela in tabelas_division.items():
        num_pessoas = len(tabela)
        if num_pessoas == 0:
            continue

        valid_counts = [c for c in df_dist_grupos.columns if pd.notna(c)]
        valid_counts = sorted(int(c) for c in valid_counts)
        lookup_count = num_pessoas
        if lookup_count not in valid_counts:
            lower_counts = [c for c in valid_counts if c <= lookup_count]
            if lower_counts:
                lookup_count = lower_counts[-1]
            else:
                lookup_count = valid_counts[0]

        max_rank = int(df_dist_grupos.index.max())
        valores_grupo = []
        for _, row in tabela.iterrows():
            rank = int(row['Rank']) if pd.notna(row['Rank']) else None
            if rank is not None:
                if rank in df_dist_grupos.index:
                    valor_grupo = df_dist_grupos.at[rank, lookup_count]
                    valores_grupo.append(int(valor_grupo))
                elif rank > max_rank:
                    valores_grupo.append(math.ceil(rank / 4))
                else:
                    valores_grupo.append(99)
            else:
                valores_grupo.append(99)

        tabela['grupo'] = valores_grupo
        tabela['pos_grupo'] = tabela.groupby('grupo').cumcount() + 1

        eliminados = tabela[tabela['grupo'] == 99]
        tabela_eliminados = pd.concat([tabela_eliminados, eliminados], ignore_index=True)

        tabela = tabela[tabela['grupo'] != 99]
        if tabela.empty:
            continue

        grupos_faltantes = tabela['grupo'].value_counts()
        grupos_completar = grupos_faltantes[grupos_faltantes < 4].index

        for grupo in grupos_completar:
            num_linhas_grupo = len(tabela[tabela['grupo'] == grupo])
            num_linhas_preencher = 4 - num_linhas_grupo

            for i in range(num_linhas_preencher):
                posicao_grupo = num_linhas_grupo + 1 + i
                linha_preenchimento = {
                    'WaID': 0,
                    'Rank': 0,
                    'Nome Completo': f'BYE {grupo}{posicao_grupo}',
                    'Clube': f'BYE {grupo}{posicao_grupo}',
                    'Sigla': '',
                    'Score': 0,
                    '10+X': 0,
                    'X': 0,
                    'Categoria Quali': tabela['Categoria Quali'].iloc[0],
                    'Categoria Combates': categoria_combate,
                    'grupo': grupo,
                    'pos_grupo': posicao_grupo,
                }
                tabela = pd.concat([tabela, pd.DataFrame([linha_preenchimento])], ignore_index=True)

        tabela.sort_values(by=['grupo', 'pos_grupo'], inplace=True)
        tabela.reset_index(drop=True, inplace=True)
        final_grupos_tabelas[categoria_combate] = tabela

        combates_divisao_df = pd.DataFrame()
        for _, grupo_df in tabela.groupby('grupo'):
            grupo_df = grupo_df.reset_index(drop=True)
            combates_data = []

            combates_data.append({
                'match': 'MATCH 1',
                'genero': categoria_combate,
                'GRUPO': f'GRUPO {grupo_df.iloc[0]["grupo"]}',
                'nome a': grupo_df.iloc[0]['Nome Completo'],
                'clube a': grupo_df.iloc[0]['Clube'],
                'rank a': grupo_df.iloc[0]['Rank'],
                'nome b': grupo_df.iloc[3]['Nome Completo'],
                'clube b': grupo_df.iloc[3]['Clube'],
                'rank b': grupo_df.iloc[3]['Rank'],
                'ETAPA': etapa,
                'LOCAL': local_da_prova,
            })
            combates_data.append({
                'match': 'MATCH 1',
                'genero': categoria_combate,
                'GRUPO': f'GRUPO {grupo_df.iloc[1]["grupo"]}',
                'nome a': grupo_df.iloc[1]['Nome Completo'],
                'clube a': grupo_df.iloc[1]['Clube'],
                'rank a': grupo_df.iloc[1]['Rank'],
                'nome b': grupo_df.iloc[2]['Nome Completo'],
                'clube b': grupo_df.iloc[2]['Clube'],
                'rank b': grupo_df.iloc[2]['Rank'],
                'ETAPA': etapa,
                'LOCAL': local_da_prova,
            })
            combates_data.append({
                'match': 'MATCH 2',
                'genero': categoria_combate,
                'GRUPO': f'GRUPO {grupo_df.iloc[0]["grupo"]}',
                'nome a': grupo_df.iloc[0]['Nome Completo'],
                'clube a': grupo_df.iloc[0]['Clube'],
                'rank a': grupo_df.iloc[0]['Rank'],
                'nome b': grupo_df.iloc[2]['Nome Completo'],
                'clube b': grupo_df.iloc[2]['Clube'],
                'rank b': grupo_df.iloc[2]['Rank'],
                'ETAPA': etapa,
                'LOCAL': local_da_prova,
            })
            combates_data.append({
                'match': 'MATCH 2',
                'genero': categoria_combate,
                'GRUPO': f'GRUPO {grupo_df.iloc[1]["grupo"]}',
                'nome a': grupo_df.iloc[1]['Nome Completo'],
                'clube a': grupo_df.iloc[1]['Clube'],
                'rank a': grupo_df.iloc[1]['Rank'],
                'nome b': grupo_df.iloc[3]['Nome Completo'],
                'clube b': grupo_df.iloc[3]['Clube'],
                'rank b': grupo_df.iloc[3]['Rank'],
                'ETAPA': etapa,
                'LOCAL': local_da_prova,
            })
            combates_data.append({
                'match': 'MATCH 3',
                'genero': categoria_combate,
                'GRUPO': f'GRUPO {grupo_df.iloc[0]["grupo"]}',
                'nome a': grupo_df.iloc[0]['Nome Completo'],
                'clube a': grupo_df.iloc[0]['Clube'],
                'rank a': grupo_df.iloc[0]['Rank'],
                'nome b': grupo_df.iloc[1]['Nome Completo'],
                'clube b': grupo_df.iloc[1]['Clube'],
                'rank b': grupo_df.iloc[1]['Rank'],
                'ETAPA': etapa,
                'LOCAL': local_da_prova,
            })
            combates_data.append({
                'match': 'MATCH 3',
                'genero': categoria_combate,
                'GRUPO': f'GRUPO {grupo_df.iloc[2]["grupo"]}',
                'nome a': grupo_df.iloc[0]['Nome Completo'],
                'clube a': grupo_df.iloc[0]['Clube'],
                'rank a': grupo_df.iloc[0]['Rank'],
                'nome b': grupo_df.iloc[1]['Nome Completo'],
                'clube b': grupo_df.iloc[3]['Clube'],
                'rank b': grupo_df.iloc[3]['Rank'],
                'ETAPA': etapa,
                'LOCAL': local_da_prova,
            })

            combates_divisao_df = pd.concat([combates_divisao_df, pd.DataFrame(combates_data)], ignore_index=True)

        combates_geral_df = pd.concat([combates_geral_df, combates_divisao_df], ignore_index=True)

    combates_para_salvar = {}
    if not combates_geral_df.empty:
        combates_sem_bye = combates_geral_df[~(
            combates_geral_df['nome a'].str.contains('BYE', na=False)
            & combates_geral_df['nome b'].str.contains('BYE', na=False)
        )]
        combates_para_salvar['Total'] = combates_sem_bye
        for categoria in combates_sem_bye['genero'].unique():
            combates_para_salvar[categoria] = combates_sem_bye[combates_sem_bye['genero'] == categoria]

    return final_grupos_tabelas, combates_para_salvar, tabela_eliminados


def _generate_docx_from_template(df: pd.DataFrame, modelo_path: str, out_docx_path: str, categoria_name: str):
    """Create a docx file for a category by copying a template and appending a table with the dataframe rows."""
    shutil.copy(modelo_path, out_docx_path)
    doc = Document(out_docx_path)
    # Add a heading with category name
    doc.add_heading(str(categoria_name), level=2)
    if df.empty:
        doc.add_paragraph('Sem combates.')
        doc.save(out_docx_path)
        return out_docx_path

    # Add a table with the dataframe contents
    cols = list(df.columns)
    table = doc.add_table(rows=1, cols=len(cols))
    hdr_cells = table.rows[0].cells
    for i, c in enumerate(cols):
        hdr_cells[i].text = str(c)

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, c in enumerate(cols):
            val = row[c]
            cells[i].text = '' if pd.isna(val) else str(val)

    doc.save(out_docx_path)
    return out_docx_path


def _convert_docx_to_pdf(docx_path: str, pdf_path: str):
    """Convert a DOCX file to PDF using docx2pdf (Windows / Word)."""
    if docx2pdf_convert is None:
        raise RuntimeError('docx2pdf is not available. Please `pip install docx2pdf` and ensure Word is installed on Windows.')
    # docx2pdf.convert accepts (input, output) where input can be file and output a folder or file
    docx2pdf_convert(docx_path, pdf_path)
    return pdf_path


def generate_print_pdf_from_combates(combates_dict: dict, output_pdf_path: str, etapa_name: str):
    """Generate per-category DOCX from templates and convert+merge into a single PDF.

    combates_dict: dict of DataFrames keyed by category name (genero).
    output_pdf_path: destination PDF full path.
    etapa_name: used for output filenames.
    Returns path to merged PDF.
    """
    base_dir = Path(__file__).resolve().parents[1]
    template_soma = base_dir / 'Modelo Robin Round Individual - Soma.docx'
    template_sets = base_dir / 'Modelo Robin Round Individual - Sets.docx'

    if not template_soma.exists() or not template_sets.exists():
        raise FileNotFoundError('Um ou ambos os modelos DOCX não foram encontrados no diretório do projeto.')

    tmpdir = Path(tempfile.mkdtemp(prefix='robinround_print_'))
    pdf_paths = []
    try:
        for genero, df in combates_dict.items():
            if genero == 'Total':
                continue
            modelo_path = template_soma if str(genero).strip().upper().startswith('C') else template_sets
            safe_name = "".join(c for c in str(genero) if c.isalnum() or c in (' ', '_', '-')).strip().replace(' ', '_')
            out_docx = tmpdir / f"{etapa_name}_{safe_name}.docx"
            _generate_docx_from_template(df, str(modelo_path), str(out_docx), genero)
            out_pdf = tmpdir / f"{etapa_name}_{safe_name}.pdf"
            _convert_docx_to_pdf(str(out_docx), str(out_pdf))
            pdf_paths.append(str(out_pdf))

        # Merge PDFs
        merger = PdfMerger()
        for p in pdf_paths:
            merger.append(p)
        with open(output_pdf_path, 'wb') as f:
            merger.write(f)
        merger.close()
        return output_pdf_path
    finally:
        # keep temp files for debugging; could remove them
        shutil.rmtree(tmpdir, ignore_errors=True)


def show_page():
    # --- Streamlit UI ---
    st.title('Gerador de Combates - Robin Round Individual')

    st.markdown('''
    Esta ferramenta automatiza a criação de planilhas de combates a partir de um arquivo de resultados brutos.

    **Instruções:**
    1. Preencha o nome da **Etapa** e o **Local da Prova**.
    2. Faça o upload do arquivo de **Resultados da Prova** (.txt, .csv ou .xlsx).
    3. Selecione quais atletas seguem para a etapa de combates.
    4. Clique em **Processar Combates**.
    5. Aguarde o processamento e faça o download dos arquivos gerados.
    ''')

    st.header('1. Informações da Prova')
    default_etapa_text = 'Xº Outdoor FPAF 2026 - Xº Robin Round Individual'
    etapa_input = st.text_input('Nome da Etapa', value=default_etapa_text)
    local_input = st.text_input('Local da Prova', '')

    st.header('2. Upload de Arquivos')
    uploaded_file = st.file_uploader('Carregue o arquivo de Resultados da Prova (.txt, .csv ou .xlsx)', type=['txt', 'csv', 'xlsx'])

    if etapa_input == default_etapa_text or not etapa_input.strip():
        st.warning('Altere o nome da etapa para um valor real antes de gerar os combates.')
        return
    if not local_input.strip():
        st.warning('Preencha o local da prova antes de gerar os combates.')
        return
    if uploaded_file is None:
        st.warning('Por favor, faça o upload do arquivo de resultados antes de gerar os combates.')
        return

    try:
        with st.spinner('Carregando resultados...'):
            df_prova = load_input_df(uploaded_file)
            df_categoria_map = load_categoria_map()
            df_prova = normalize_input(df_prova, df_categoria_map)

        if df_prova.empty:
            st.warning('O arquivo carregado não contém atletas válidos.')
            return

        st.header('3. Seleção de Atletas')
        st.write('Abaixo estão os atletas carregados. Desmarque os que NÃO avançam para a etapa de combates.')

        df_prova['Target'] = df_prova.get('Target', '').fillna('').astype(str)
        df_prova['Nome Completo'] = df_prova['Nome Completo'].fillna('').astype(str)
        df_prova['Categoria Quali'] = df_prova['Categoria Quali'].fillna('').astype(str)
        df_prova['Sigla'] = df_prova['Sigla'].fillna('').astype(str)

        selection_df = df_prova[['Target', 'Nome Completo', 'Categoria Quali', 'Sigla']].copy().reset_index(drop=False)
        selection_df = selection_df.rename(columns={'Target': 'Alvo'})

        if 'athlete_selection_confirmed' not in st.session_state:
            st.session_state['athlete_selection_confirmed'] = False

        with st.form('athlete_selection_form'):
            header_cols = st.columns([1.0, 4.0, 2.0, 1.2, 2.0])
            header_cols[0].write('Combate')
            header_cols[1].write('Nome Completo')
            header_cols[2].write('Categoria Quali')
            header_cols[3].write('Sigla')
            header_cols[4].write('Alvo')

            for _, row in selection_df.iterrows():
                cols = st.columns([1.0, 4.0, 2.0, 1.2, 2.0])
                cols[0].checkbox('', value=True, key=f'sel_{uploaded_file.name}_{row["index"]}')
                cols[1].write(row['Nome Completo'])
                cols[2].write(row['Categoria Quali'])
                cols[3].write(row['Sigla'])
                cols[4].write(row['Alvo'])

            submitted = st.form_submit_button('Confirmar seleção')

        if submitted:
            st.session_state['athlete_selection_confirmed'] = True

        if not st.session_state['athlete_selection_confirmed']:
            return

        selected_indexes = []
        for _, row in selection_df.iterrows():
            if st.session_state.get(f'sel_{uploaded_file.name}_{row["index"]}', False):
                selected_indexes.append(row['index'])

        if not selected_indexes:
            st.warning('Marque pelo menos um atleta como selecionado para avançar.')
            return

        st.info(f'{len(selected_indexes)} atleta(s) selecionado(s) para a etapa de combates.')
        if not st.button('Processar Combates', key='processar_combates_button'):
            return

        with st.spinner('Processando... Por favor, aguarde.'):
            df_selecionados = df_prova.loc[selected_indexes].copy()
            df_selecionados = calculate_rank(df_selecionados)
            df_dist_grupos = load_dist_grupos()
            grupos_final, combates_final, eliminados_final = processar_combates(
                df_selecionados, df_dist_grupos, etapa_input, local_input
            )

        st.success('Processamento concluído com sucesso!')
        st.session_state['athlete_selection_confirmed'] = False

        if grupos_final:
            grupos_xlsx = to_excel(grupos_final, multi_sheet=True)
            st.download_button(
                label='⬇️ Baixar Planilha de Grupos',
                data=grupos_xlsx,
                file_name=f'{etapa_input}_grupos.xlsx',
                mime='application/vnd.openxmlformats-officedocument-spreadsheetml.sheet',
            )

        if combates_final:
            combates_xlsx = to_excel(combates_final, multi_sheet=True)
            st.download_button(
                label='⬇️ Baixar Planilha de Combates',
                data=combates_xlsx,
                file_name=f'{etapa_input}_combates.xlsx',
                mime='application/vnd.openxmlformats-officedocument-spreadsheetml.sheet',
            )

            st.markdown('**Impressão:** Gere um PDF único com as páginas por categoria a partir dos modelos.')
            if docx2pdf_convert is None:
                st.warning('Conversão DOCX→PDF indisponível: instale `docx2pdf` e tenha o Word no Windows para gerar PDFs automaticamente.')
            generate_pdf = st.button('Gerar PDFs para Impressão', key='gerar_pdfs_button')
            if generate_pdf:
                try:
                    tmp_out_pdf = Path(tempfile.gettempdir()) / f"{etapa_input}_impressao.pdf"
                    with st.spinner('Gerando arquivos DOCX e convertendo para PDF (pode demorar)...'):
                        pdf_path = generate_print_pdf_from_combates(combates_final, str(tmp_out_pdf), etapa_input)
                    with open(pdf_path, 'rb') as f:
                        pdf_bytes = f.read()
                    st.success('PDF de impressão gerado com sucesso.')
                    st.download_button('⬇️ Baixar PDF para Impressão', data=pdf_bytes, file_name=f'{etapa_input}_impressao.pdf', mime='application/pdf')
                except Exception as e:
                    st.error(f'Falha ao gerar o PDF de impressão: {e}')

        if not eliminados_final.empty:
            eliminados_xlsx = to_excel(eliminados_final)
            st.download_button(
                label='⬇️ Baixar Planilha de Eliminados',
                data=eliminados_xlsx,
                file_name=f'{etapa_input}_eliminados.xlsx',
                mime='application/vnd.openxmlformats-officedocument-spreadsheetml.sheet',
            )
        else:
            st.info('Não há eliminados.')

    except FileNotFoundError as e:
        st.error(str(e))
    except Exception as e:
        st.error(f'Ocorreu um erro durante o processamento: {e}')
        st.warning('Verifique se o arquivo de resultados está no formato correto e tente novamente.')
